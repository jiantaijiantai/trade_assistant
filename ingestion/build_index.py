\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\
\


import argparse
import hashlib
import os
from pathlib import Path

import pdfplumber
import pytesseract
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image

from rag.access_control import (
    BUSINESS_COLLECTION_NAME,
    BUSINESS_DEPARTMENT_ID,
    DEFAULT_TENANT_ID,
    build_default_acl_metadata,
)
from rag.chroma_store import (
    DEFAULT_CHROMA_DIR,
    DEFAULT_COLLECTION_NAME,
    delete_vectors,
    upsert_vectors,
)
from rag.chunker import build_file_id, chunk_document
from rag.embeddings import embed_texts, load_embedding_config
from rag.lexical_index import (
    DEFAULT_LEXICAL_INDEX_PATH,
    delete_lexical_chunks,
    upsert_lexical_chunks,
)
from rag.schemas import IndexManifest, IndexedFileRecord, ParsedDocument


DEFAULT_MANIFEST_PATH = "outputs/rag/business_index_manifest.json"
SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf", ".xlsx", ".jpg", ".jpeg", ".png"}


def sha256_file(path: Path) -> str:
\
\
\
\
\
\
\


    digest = hashlib.sha256()

    with path.open("rb") as file:
        for block in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(block)

    return digest.hexdigest()


def scan_files(input_dir: Path) -> list[Path]:
\
\
\
\
\
\
\
\
\
\


    files: list[Path] = []

    for path in input_dir.rglob("*"):
        if not path.is_file():
            continue

        if path.name.startswith("~$"):
            continue

        if path.name.startswith("."):
            continue

        if path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(path)

    return sorted(files)


def read_txt(path: Path) -> tuple[str, dict, list[str]]:
\
\
\
\
\


    warnings: list[str] = []

    for encoding in ("utf-8", "gbk"):
        try:
            return path.read_text(encoding=encoding), {"encoding": encoding}, warnings
        except UnicodeDecodeError:
            continue

    warnings.append("TXT 编码无法识别，已忽略正文")
    return "", {}, warnings


def read_docx(path: Path) -> tuple[str, dict, list[str]]:
\
\
\
\
\
\
\
\


    warnings: list[str] = []
    document = DocxDocument(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    table_count = 0

    for table in document.tables:
        table_count += 1

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts), {"table_count": table_count}, warnings


def read_pdf(path: Path) -> tuple[str, dict, list[str]]:
\
\
\
\
\
\


    warnings: list[str] = []
    parts: list[str] = []
    page_count = 0

    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)

        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()

            if text:
                parts.append(f"[PDF 第 {index} 页]\n{text}")

    if not parts:
        warnings.append("PDF 未提取到文本层，可能是扫描件")

    return "\n\n".join(parts), {"page_count": page_count}, warnings


def read_xlsx(path: Path) -> tuple[str, dict, list[str]]:
\
\
\
\
\
\
\
\


    warnings: list[str] = []
    workbook = load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []

    for sheet in workbook.worksheets:
        non_empty_rows = 0

        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if not values:
                continue

            non_empty_rows += 1
            parts.append(f"[Sheet:{sheet.title} Row:{row_index}] " + " | ".join(values))

        if non_empty_rows == 0:
            warnings.append(f"Excel sheet 为空：{sheet.title}")

    return "\n".join(parts), {"sheet_count": len(workbook.worksheets)}, warnings


def read_image(
    path: Path,
    tesseract_cmd: str,
    tessdata_dir: str,
    ocr_lang: str,
) -> tuple[str, dict, list[str]]:
\
\
\
\
\
\
\
\
\


    warnings: list[str] = []

    if not tesseract_cmd:
        warnings.append("未配置 tesseract_cmd，图片未 OCR")
        return "", {}, warnings

    if not Path(tesseract_cmd).exists():
        warnings.append(f"tesseract_cmd 不存在：{tesseract_cmd}")
        return "", {}, warnings

    if tessdata_dir and not Path(tessdata_dir).exists():
        warnings.append(f"tessdata_dir 不存在：{tessdata_dir}")
        return "", {}, warnings

    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    try:
        image = Image.open(path)
        config = f'--tessdata-dir "{tessdata_dir}"' if tessdata_dir else ""
        text = pytesseract.image_to_string(image, lang=ocr_lang, config=config)
        text = text.strip()

        if not text:
            warnings.append("图片 OCR 未识别到文本")

        return text, {"ocr_lang": ocr_lang}, warnings

    except Exception as exc:
        warnings.append(f"图片 OCR 失败：{exc}")
        return "", {"ocr_lang": ocr_lang}, warnings


def parse_file(
    path: Path,
    tesseract_cmd: str,
    tessdata_dir: str,
    ocr_lang: str,
) -> ParsedDocument:
\
\
\
\
\


    file_hash = sha256_file(path)
    file_type = path.suffix.lower()

    if file_type == ".txt":
        text, metadata, warnings = read_txt(path)
    elif file_type == ".docx":
        text, metadata, warnings = read_docx(path)
    elif file_type == ".pdf":
        text, metadata, warnings = read_pdf(path)
    elif file_type == ".xlsx":
        text, metadata, warnings = read_xlsx(path)
    elif file_type in {".jpg", ".jpeg", ".png"}:
        text, metadata, warnings = read_image(
            path=path,
            tesseract_cmd=tesseract_cmd,
            tessdata_dir=tessdata_dir,
            ocr_lang=ocr_lang,
        )
    else:
        text, metadata, warnings = "", {}, [f"不支持的文件类型：{file_type}"]

    if not text.strip():
        warnings.append("未解析出可入库文本")

    return ParsedDocument(
        source_path=str(path),
        file_name=path.name,
        file_type=file_type,
        file_hash=file_hash,
        text=text,
        metadata=metadata,
        warnings=warnings,
    )


def load_manifest(
    manifest_path: Path,
    embedding_provider: str,
    embedding_model: str,
) -> IndexManifest:


    if not manifest_path.exists():
        return IndexManifest(
            vector_store="chroma",
            collection_name=DEFAULT_COLLECTION_NAME,
            embedding_provider=embedding_provider,
            embedding_model=embedding_model,
            files={},
        )

    return IndexManifest.model_validate_json(manifest_path.read_text(encoding="utf-8"))


def save_manifest(manifest: IndexManifest, manifest_path: Path) -> None:


    manifest_path.parent.mkdir(parents=True, exist_ok=True)

    tmp_path = manifest_path.with_suffix(manifest_path.suffix + ".tmp")
    tmp_path.write_text(
        manifest.model_dump_json(indent=2),
        encoding="utf-8",
    )
    tmp_path.replace(manifest_path)


def remove_old_file_chunks(
    record: IndexedFileRecord,
    chroma_dir: str,
    collection_name: str,
    lexical_index_path: str,
) -> None:
\
\
\
\
\
\
\
\


    delete_vectors(
        record.chunk_ids,
        persist_dir=chroma_dir,
        collection_name=collection_name,
    )
    delete_lexical_chunks(
        record.chunk_ids,
        index_path=lexical_index_path,
    )

def index_document(
    document: ParsedDocument,
    manifest: IndexManifest,
    chunk_size: int,
    overlap: int,
    chroma_dir: str,
    collection_name: str,
    lexical_index_path: str,
) -> None:
\
\
\
\
\
\
\
\
\
\


    chunks = chunk_document(
        document=document,
        chunk_size=chunk_size,
        overlap=overlap,
    )

    file_id = build_file_id(document.source_path, document.file_hash)

    if not chunks:
        manifest.files[document.source_path] = IndexedFileRecord(
            file_id=file_id,
            source_path=document.source_path,
            file_name=document.file_name,
            file_type=document.file_type,
            file_hash=document.file_hash,
            chunk_ids=[],
            last_indexed_at=os.getenv("CURRENT_INDEX_TIME", ""),
            warnings=document.warnings,
        )
        return

    embeddings = embed_texts([chunk.text for chunk in chunks])

    upsert_vectors(
        chunks,
        embeddings,
        persist_dir=chroma_dir,
        collection_name=collection_name,
    )
    upsert_lexical_chunks(
        chunks,
        index_path=lexical_index_path,
    )

    manifest.files[document.source_path] = IndexedFileRecord(
        file_id=file_id,
        source_path=document.source_path,
        file_name=document.file_name,
        file_type=document.file_type,
        file_hash=document.file_hash,
        chunk_ids=[chunk.chunk_id for chunk in chunks],
        last_indexed_at=os.getenv("CURRENT_INDEX_TIME", ""),
        warnings=document.warnings,
    )

def build_index(args: argparse.Namespace) -> dict[str, int]:
\
\
\
\


    input_dir = Path(args.input)
    manifest_path = Path(args.manifest)

    if not input_dir.exists():
        raise FileNotFoundError(f"输入目录不存在：{input_dir}")

    embedding_config = load_embedding_config()

    manifest = load_manifest(
        manifest_path=manifest_path,
        embedding_provider=embedding_config.provider,
        embedding_model=embedding_config.model,
    )

    files = scan_files(input_dir)
    current_paths = {str(path) for path in files}

    stats = {
        "scanned_files": len(files),
        "indexed_files": 0,
        "skipped_files": 0,
        "changed_files": 0,
        "deleted_files": 0,
        "failed_files": 0,
    }

    if args.mode == "rebuild":
        for record in list(manifest.files.values()):
            remove_old_file_chunks(
                record,
                chroma_dir=args.chroma_dir,
                collection_name=args.collection_name,
                lexical_index_path=args.lexical_index,
            )
        manifest.files = {}

    if args.mode == "sync":
        for source_path, record in list(manifest.files.items()):
            if source_path not in current_paths:
                remove_old_file_chunks(
                    record,
                    chroma_dir=args.chroma_dir,
                    collection_name=args.collection_name,
                    lexical_index_path=args.lexical_index,
                )
                manifest.files.pop(source_path, None)
                stats["deleted_files"] += 1

    for path in files:
        source_path = str(path)
        file_hash = sha256_file(path)
        existing = manifest.files.get(source_path)

        if args.mode in {"sync", "append"} and existing and existing.file_hash == file_hash:
            stats["skipped_files"] += 1
            continue

        if existing and existing.file_hash != file_hash:
            remove_old_file_chunks(
                existing,
                chroma_dir=args.chroma_dir,
                collection_name=args.collection_name,
                lexical_index_path=args.lexical_index,
            )
            stats["changed_files"] += 1

        try:
            document = parse_file(
                path=path,
                tesseract_cmd=args.tesseract_cmd,
                tessdata_dir=args.tessdata_dir,
                ocr_lang=args.ocr_lang,
            )
            document.metadata.update(
                build_default_acl_metadata(
                    tenant_id=getattr(args, "tenant_id", DEFAULT_TENANT_ID),
                    department_id=getattr(args, "department_id", BUSINESS_DEPARTMENT_ID),
                    owner_user_id=getattr(args, "owner_user_id", ""),
                    visibility=getattr(args, "visibility", "department"),
                    allowed_user_ids=getattr(args, "allowed_user_ids", ""),
                    allowed_roles=getattr(args, "allowed_roles", ""),
                    allowed_groups=getattr(args, "allowed_groups", ""),
                    sensitivity_level=getattr(args, "sensitivity_level", "internal"),
                )
            )
            index_document(
                document=document,
                manifest=manifest,
                chunk_size=args.chunk_size,
                overlap=args.overlap,
                chroma_dir=args.chroma_dir,
                collection_name=args.collection_name,
                lexical_index_path=args.lexical_index,
            )
            stats["indexed_files"] += 1
        except Exception as exc:
            stats["failed_files"] += 1
            print(f"[ERROR] 文件入库失败：{path} -> {exc}")

    manifest.embedding_provider = embedding_config.provider
    manifest.embedding_model = embedding_config.model
    manifest.collection_name = args.collection_name

    save_manifest(manifest, manifest_path)

    return stats


def parse_args() -> argparse.Namespace:


    parser = argparse.ArgumentParser(description="构建 trade_assistant 业务资料 RAG 索引")

    parser.add_argument("--input", required=True, help="业务资料目录，例如：业务资料")
    parser.add_argument(
        "--mode",
        choices=["rebuild", "sync", "append"],
        default="sync",
        help="索引模式：rebuild / sync / append，默认 sync",
    )
    parser.add_argument(
        "--manifest",
        default=DEFAULT_MANIFEST_PATH,
        help="manifest 保存路径",
    )
    parser.add_argument(
        "--chroma-dir",
        default=DEFAULT_CHROMA_DIR,
        help="Chroma 持久化目录",
    )
    parser.add_argument(
        "--collection-name",
        default=BUSINESS_COLLECTION_NAME,
        help="Chroma collection 名称",
    )
    parser.add_argument(
        "--lexical-index",
        default=DEFAULT_LEXICAL_INDEX_PATH,
        help="BM25/关键词索引路径",
    )

    parser.add_argument("--chunk-size", type=int, default=800, help="chunk 最大字符数")
    parser.add_argument("--overlap", type=int, default=120, help="相邻 chunk 重叠字符数")
    parser.add_argument("--tenant-id", default=DEFAULT_TENANT_ID, help="RAG 租户 ID")
    parser.add_argument("--department-id", default=BUSINESS_DEPARTMENT_ID, help="RAG 部门 collection key")
    parser.add_argument("--owner-user-id", default="", help="资料上传人/所有人 ID，仅用于审计")
    parser.add_argument(
        "--visibility",
        default="department",
        choices=["department", "group", "role", "users"],
        help="资料可见范围",
    )
    parser.add_argument("--allowed-user-ids", default="", help="逗号分隔的员工白名单")
    parser.add_argument("--allowed-roles", default="", help="逗号分隔的角色白名单")
    parser.add_argument("--allowed-groups", default="", help="逗号分隔的权限组白名单")
    parser.add_argument(
        "--sensitivity-level",
        default="internal",
        choices=["internal", "confidential", "restricted"],
        help="资料密级",
    )
    parser.add_argument(
        "--tesseract-cmd",
        default=r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        help="tesseract.exe 路径",
    )
    parser.add_argument(
        "--tessdata-dir",
        default=r"C:\Users\86182\Desktop\分阶段双项目\ocr_tessdata",
        help="Tesseract 语言包目录",
    )
    parser.add_argument("--ocr-lang", default="chi_sim+eng", help="OCR 语言")

    return parser.parse_args()


def main() -> None:


    args = parse_args()
    stats = build_index(args)

    print("阶段 3 RAG 索引构建完成")
    for key, value in stats.items():
        print(f"{key}: {value}")


if __name__ == "__main__":
    main()
